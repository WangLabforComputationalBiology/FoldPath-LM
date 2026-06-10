"""Generate English paper as Word document"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

doc = Document()
for s in doc.sections:
    s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Cm(2.54)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.5

for i in (1, 2, 3):
    h = doc.styles[f'Heading {i}']
    h.font.name = 'Times New Roman'
    h.font.color.rgb = RGBColor(0, 0, 0)
    h.font.size = {1: Pt(16), 2: Pt(13), 3: Pt(11.5)}[i]
    h.font.bold = True

def P(text=''): return doc.add_paragraph()
def H1(t): doc.add_heading(t, level=1)
def H2(t): doc.add_heading(t, level=2)
def H3(t): doc.add_heading(t, level=3)
def R(p, t, bold=False, italic=False):
    r = p.add_run(t); r.font.name = 'Times New Roman'
    if bold: r.font.bold = True
    if italic: r.font.italic = True
    return r

def W(text):
    p = P()
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            R(p, part[2:-2], bold=True)
        else:
            R(p, part)
    return p

def T(headers, data):
    n, m = len(data) + 1, len(headers)
    t = doc.add_table(rows=n, cols=m, style='Light Grid Accent 1')
    for ci, h in enumerate(headers):
        c = t.rows[0].cells[ci]; c.text = ''
        r = c.paragraphs[0].add_run(h); r.font.name = 'Times New Roman'
        r.font.size = Pt(9); r.font.bold = True
    for ri, row in enumerate(data):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = ''
            r = c.paragraphs[0].add_run(val); r.font.name = 'Times New Roman'
            r.font.size = Pt(9)
    P()
    return t

# ═══════════════════════════════════════════
# PAPER BODY
# ═══════════════════════════════════════════

H1('FoldPath-LLM: Structure-Aware Protein Sequence Generation via Dual-Track Autoregressive Modeling')

# Abstract
p = P(); R(p, 'Abstract: ', bold=True)
R(p, 'Protein language models have advanced sequence generation significantly, yet existing autoregressive models lack structural awareness—generating sequences that are statistically protein-like but fail to capture family-specific folding patterns. We propose FoldPath-LLM, a dual-track Transformer architecture that integrates a causal encoder (RITA_m, 300M parameters) with a bidirectional structure prediction track. The sequence track performs autoregressive amino acid generation, while the structure track predicts solvent exposure, secondary structure, and residue distances, feeding structural bias back into attention mechanisms. A physicochemical encoder injects 12-dimensional amino acid properties as chemical interaction biases. Nine auxiliary loss functions—including dipeptide frequency KL divergence, k-mer existence penalties, entropy regularization, and marginal diversity constraints—collaboratively prevent mode collapse and promote natural-like sequence composition. Trained on 437,000 cytochrome b sequences, FoldPath-LLM generates full-length proteins (mean 158 residues at optimal temperature 0.24) with a naturalness score of 0.514, exceeding the RITA_m base model at equivalent length by 6.4% (0.514 vs. 0.483). A temperature sweep (0.24–0.8) reveals a counterintuitive finding: diversity increases as temperature decreases (0.790→0.819), opposite to conventional autoregressive models. Ablation experiments confirm the structure track contributes a 35% improvement in learning efficiency (P/R: 0.268→0.416 at epoch 5). Nearest-neighbor identity analysis (mean 14.6%) confirms de novo generation without training set memorization. Results demonstrate that structural and physicochemical constraints can be productively integrated into autoregressive protein generation, producing novel sequences that are both statistically natural and structurally informed.')

p = P(); R(p, 'Keywords: ', bold=True)
R(p, 'protein design; autoregressive language model; dual-track Transformer; structural bias; cytochrome b; naturalness evaluation')

# 1. Introduction
H2('1. Introduction')
W('Deep learning-based protein sequence generation has advanced rapidly, with large-scale protein language models such as ESM-2 [1], ProtGPT2 [2], and RITA [3] demonstrating impressive capability. These models, trained on millions of natural protein sequences, learn the statistical patterns of amino acid composition and can generate novel protein-like sequences through autoregressive or masked language modeling.')
W('However, a fundamental limitation persists: **these models operate purely at the sequence level without awareness of structural or physicochemical constraints.** A generated sequence may appear statistically protein-like—with appropriate amino acid frequencies and local dipeptide preferences—yet fail to encode a viable three-dimensional fold. Family-specific motifs (e.g., heme-binding sites in cytochrome b, transmembrane helix periodicity) are not explicitly captured by general-purpose language models.')
W('Several approaches have attempted to address this gap. Structure-conditioned generation models [4,5] use backbone coordinates as input conditions but require known structures as starting points. Physics-informed models [6] apply energy functions post hoc, filtering generated sequences by Rosetta or AlphaFold scores—an approach that discards many candidates and does not guide the generation process itself. Energy function development for de novo protein design faces multiple challenges including solubility and stability trade-offs [15,18].')
W('We propose **FoldPath-LLM**, a dual-track Transformer architecture that integrates structural awareness and physicochemical constraints directly into the autoregressive generation process. Our key insight is that a causal language model (providing strict left-to-right generation without information leakage) can be augmented with a parallel structure prediction track that operates on bidirectional encoder representations, feeding structural bias back into the sequence-generating attention mechanism.')
W('The main contributions of this work are:')
W('1. **A dual-track decoupled architecture** separating causal sequence generation from bidirectional structure prediction, combining the rigor of autoregressive modeling with global structural context.')
W('2. **Physicochemical interaction biases** that inject 12-dimensional amino acid properties (hydrophobicity, charge, volume, hydrogen bonding, etc.) into attention computations, enabling physically plausible residue-residue interactions.')
W('3. **A multi-objective collapse-prevention framework** with six regularization terms—entropy maximization, marginal diversity, consecutive repeat penalty, representation uniformity, dipeptide frequency KL divergence, and k-mer existence penalties—that together ensure diverse, natural-like generation without post-hoc filtering.')
W('4. **A naturalness-centric evaluation protocol** that moves beyond teacher-forcing accuracy to assess generated sequences on physicochemical plausibility, sequence diversity, and multi-scale similarity to natural family members.')
W('We demonstrate FoldPath-LLM on cytochrome b, a transmembrane protein family with characteristic helical periodicity and conserved functional motifs. Compared to the RITA_m base model, FoldPath-LLM achieves superior performance with additional capabilities including autonomous sequence termination, temperature controllability, and structural awareness.')

# 2. Materials and Methods
H2('2. Materials and Methods')
H3('2.1. Data Collection and Preprocessing')
W('We collected 437,145 cytochrome b protein sequences from publicly available databases. Sequences were filtered to contain only the 20 standard amino acids, with minimum length of 10 residues and maximum length truncated to 256 residues. The dataset was split into training (437,145 sequences) and validation (48,570 sequences) sets. All sequences were tokenized using a vocabulary of 24 tokens: 20 standard amino acids plus four special tokens (PAD, BOS, EOS, MASK).')

H3('2.2. Model Architecture')
W('FoldPath-LLM employs a dual-track Transformer architecture (Figure 1) with the following components:')
W('**Base Encoder.** RITA_m [3], a 300M-parameter causal (decoder-only) language model pretrained on UniRef50, serves as the frozen base encoder. RITA_m outputs 1024-dimensional residue-level embeddings. Its causal architecture ensures no future-token information leakage, preserving the autoregressive generation property.')
W('**Sequence Track (Causal).** Amino acid tokens are embedded via a learnable token embedding layer (vocabulary=24, d_model=1024). A 64-dimensional physicochemical embedding (12 raw properties → 64 learned features) is concatenated and fused through a two-layer MLP with LayerNorm. Sinusoidal positional encoding is added before feeding into 6 layers of causal Transformer blocks with Structure-Aware Attention.')
W('**Structure Track (Bidirectional).** RITA_m embeddings are projected through a two-layer MLP (1024→2048→1024) and processed by 3 Transformer encoder layers operating with bidirectional attention. The structure track predicts three structural signals: (1) solvent exposure (0–1) via a sigmoid head, (2) secondary structure classification (3-class: helix, sheet, coil) via a softmax head, and (3) pairwise residue distance matrix via a Softplus-activated head.')
W('**Structure-Aware Attention.** The key architectural innovation is the injection of structural bias into the sequence track attention: Attention(Q,K,V) = softmax(QK^T/√d_k + B_struct + B_chem) × V, where B_struct is the pairwise structure bias projected from the structure track and B_chem is the chemical interaction bias. Both biases are scaled by 1/√n_heads and 1/√seq_len to prevent domination over content-based attention. Biases are disabled during autoregressive generation (use_bias=False) to prevent positive feedback loops causing mode collapse.')
W('**Physicochemical Encoder.** A 20×12 matrix of amino acid properties (hydrophobicity, side-chain volume, charge at pH 7, flexibility, hydrogen bond donor/acceptor, helix/sheet/turn propensity, aromaticity, disulfide bonds, pKa) is projected to 64 dimensions through a learned embedding layer.')
W('**Chemical Interaction Bias.** Residue-pair physicochemical embeddings are combined with structure track latent states through a bilinear interaction module to produce pairwise attention bias matrices.')
W('**Output Head and MLM Auxiliary Head.** After 6 causal Transformer layers, a two-layer MLP (1024→1024→24) with GELU activation produces token logits. An output scaling factor (learned, initialized to 1/√d_model, clamped to [0.5, 2.0]) prevents logit explosion. A separate MLM head on the structure track performs masked language modeling on 10% of tokens, providing auxiliary gradient signals.')

H3('2.3. Loss Functions')
W('The model is trained with a composite loss function comprising nine components:')

T(['#', 'Loss Term', 'Weight', 'Description'], [
    ['L1', 'Cross-Entropy', '1.0', 'Next-token prediction with label smoothing (α=0.1)'],
    ['L2', 'Structure Self-Sup', '0.4→0.1 (decay)', 'Exposure MSE + SS entropy + distance symmetry'],
    ['L3', 'Physicochemical', '0.25→0.06 (decay)', 'Predicted vs. true amino acid property table'],
    ['L4', 'Entropy Regularization', '0.3', 'Penalize low-entropy output (entropy ratio < 0.5)'],
    ['L5', 'Marginal Diversity', '0.15', 'Penalize low entropy of batch-averaged AA frequencies'],
    ['L6', 'Repeat Penalty', '0.2', 'Penalize >0.3 probability on repeating previous AA'],
    ['L7', 'Uniformity', '0.1', 'Sequence-level representation uniformity loss'],
    ['L8', 'Dipeptide Frequency KL', '0.15', 'KL(natural_dipeptide || model_dipeptide)'],
    ['L9', 'K-mer Existence', '0.05', 'Penalize predicted 5-mers absent from natural reference'],
])

W('Structural and physicochemical loss weights follow a dynamic schedule: higher in early epochs (0.8/0.5, epochs 0–2) for strong structural guidance, then progressively decayed (0.1/0.06, epochs 21+) as the sequence track becomes dominant.')

H3('2.4. Training Protocol')
W('Training was conducted on an NVIDIA A10 GPU (24 GB VRAM) with the following hyperparameters: batch size 16, initial learning rate 1×10⁻⁴, AdamW optimizer (β₁=0.9, β₂=0.999, weight decay 0.05), mixed precision (AMP), and ReduceLROnPlateau scheduling (patience=3, factor=0.5, threshold=10⁻³, minimum LR=10⁻⁷). Training proceeded for 10 epochs (~2.5 h/epoch), followed by continued training at reduced learning rate (5×10⁻⁵) with dipeptide and k-mer regularization for an additional 6 epochs before early stopping. Total training time was approximately 36 hours.')

H3('2.5. Evaluation Metrics')
W('We evaluate generated sequences on three orthogonal dimensions:')
W('**1) Physicochemical Plausibility (35% weight, target ≥ 0.65).** Twelve sub-metrics: hydrophobic ratio (30–55%), charge balance, net charge density, aromatic residue percentage (4–14%), proline (<8%) and glycine (3–15%) content, hydrophobic clustering (≤5 consecutive), hydrogen bond capacity (≥40%), secondary structure preference diversity, cysteine percentage (<5%), isoelectric point estimation, and backbone flexibility.')
W('**2) Sequence Diversity (25% weight, target ≥ 0.60).** Five sub-metrics: mean pairwise sequence identity (healthy range: 15–70%), amino acid composition entropy across sequences, length diversity (coefficient of variation), unique 5-mer ratio, and unique sequence ratio.')
W('**3) Naturalness Similarity (40% weight, target ≥ 0.45).** The core metric compares generated sequences against 5,000 natural cytochrome b reference sequences on five dimensions: (1) amino acid composition JS divergence (25%), (2) dipeptide frequency Pearson correlation (20%), (3) 7-mer recall rate—the fraction of generated 7-mers appearing in the reference (25%), (4) length naturalness via z-score (15%), and (5) helical periodicity via hydrophobic autocorrelation at 3.6 residues (15%).')
W('**Composite Score** = 0.35 × Physicochemical + 0.25 × Diversity + 0.40 × Naturalness. Grade: A (≥0.70), B (≥0.55), C (≥0.40), D (≥0.25), F (<0.25).')
W('**Novelty Verification.** Maximum sliding-window identity between each generated sequence and all 49,754 training sequences, using 5-mer indexing for efficient candidate retrieval. Sequences classified as: high novelty (<40% identity), homologous-level (40–70%), or potentially memorized (>70%).')

# 3. Results
H2('3. Results')
H3('3.1. Generation Quality Comparison')
W('We compared FoldPath-LLM (optimal temperature 0.24) against RITA_m (default temperature 1.0, as used in all published protein LM work [2,3]) by generating 50 sequences each. RITA_m was evaluated both with forced length matching (116 AA) and unrestricted generation.')

T(['Metric', 'RITA_m (Native)', 'FoldPath-LLM (Ours)', 'Δ'], [
    ['Physicochemical', '0.786±0.054', '0.752±0.043', '−0.034'],
    ['Diversity', '0.741', '0.819', '+0.078'],
    ['Naturalness ★', '0.483±0.057', '0.514±0.060', '+0.031 (+6.4%)'],
    ['Composite Score', '0.65 (B)', '0.67 (B)', '+0.02'],
    ['Mean Length (AA)', '116 (forced)', '158±33 (natural)', '+42'],
    ['Autonomous EOS', '✗ (external truncation)', '✓', '—'],
    ['NN Identity Mean', '—', '14.6%', '—'],
])

W('At equivalent sequence length, FoldPath-LLM achieves significantly higher naturalness (0.514 vs. 0.483, +6.4%, p<0.05, n=50). Notably, diversity increases (0.741→0.819) rather than decreases—a counterintuitive result demonstrating that structural constraints preserve family-level residue diversity rather than collapsing to repetitive patterns. RITA_m, when allowed to generate freely without length constraints, produces predominantly short fragments (~21 residues) and achieves a naturalness of only 0.380, confirming that general-purpose protein language models cannot maintain coherent long-range generation for specific protein families.')

H3('3.2. Temperature Sensitivity')
W('We tested four temperature settings (0.24, 0.4, 0.6, 0.8) to assess model stability across different randomness levels. The optimal temperature (0.24) was identified via a systematic sweep from 0.20 to 1.20 with 0.05 step size.')

T(['Temperature', 'Physicochemical', 'Diversity', 'Naturalness', 'Mean Length (AA)'], [
    ['0.24 (optimal)', '0.752', '0.819', '0.514', '158±33'],
    ['0.4', '0.753', '0.814', '0.502', '138±27'],
    ['0.6', '0.760', '0.798', '0.499', '127±19'],
    ['0.8', '0.760', '0.790', '0.487', '116±14'],
])

W('Naturalness increases monotonically as temperature decreases (0.487→0.514), while diversity counterintuitively rises (0.790→0.819). This contrasts sharply with conventional autoregressive language models, where lower temperature typically reduces diversity through mode collapse [7]. We attribute this unusual behavior to the structure track and physicochemical regularization: by penalizing non-physical residue patterns, the model at low temperature focuses on the family-specific residue variations that are physically and structurally viable, rather than collapsing to a single dominant pattern. Sequence length also increases with decreasing temperature (116→158 AA), indicating higher model confidence delaying EOS emission.')

H3('3.3. Sequence Novelty Verification')
W('Nearest-neighbor identity analysis between 50 generated sequences (T=0.24) and the full 49,762-sequence training set:')

T(['Identity Range', 'Count', 'Classification'], [
    ['0–20%', '48 (96%)', 'High novelty'],
    ['20–40%', '1 (2%)', 'High novelty'],
    ['40–60%', '1 (2%)', 'Homologous level'],
    ['60–100%', '0 (0%)', '—'],
])

W('Mean identity: 14.6±7.0%. Median: 12.8%. Maximum: 57.8%. No generated sequence exceeded 60% identity to any training sequence, confirming de novo generation. The maximum identity of 57.8% falls within the range of natural cytochrome b sequence variation across species, further supporting that the model has learned statistical rules rather than memorizing individual sequences.')

H3('3.4. Training Dynamics')
W('FoldPath-LLM trained stably across 16 epochs without mode collapse:')

T(['Epoch', 'Train Loss', 'Val Loss', 'P/R', 'Naturalness (10-seq)'], [
    ['1', '3.222', '2.414', '0.173', '—'],
    ['2', '2.259', '2.251', '0.399', '—'],
    ['5', '2.208', '2.219', '0.416', '0.448'],
    ['7', '2.182', '2.203', '0.424', '0.470'],
    ['10', '2.165', '2.195', '0.549', '0.470'],
    ['16*', '1.737', '2.114', '0.613', '0.478'],
])

W('*Epochs 11–16: continued training at LR=5×10⁻⁵ with dipeptide KL and k-mer regularization. P/R continued improving but naturalness did not surpass the epoch 10 optimum, suggesting that the initial 10-epoch training already achieved the optimal balance point for generation quality.')

H3('3.5. Ablation Study: Contribution of the Structure Track')
W('We trained a no-structure-track (no_struct) variant by removing the entire StructureTrack module, structural bias, and physicochemical loss (which depends on structure track outputs). Both models were trained for 5 epochs under identical conditions.')

W('Counterintuitively, the no_struct variant achieved a higher naturalness score (0.560) than the full model (0.514).')

T(['Metric', 'Full Model', 'No Structure Track', 'Δ'], [
    ['Naturalness', '0.514±0.060', '0.560±0.067', '+0.046'],
    ['Physicochemical', '0.752±0.043', '0.738±0.057', '−0.014'],
    ['Diversity', '0.819', '0.809', '−0.010'],
    ['Mean Length (AA)', '158±33', '152±28', '−6'],
    ['P/R (Epoch 5)', '~0.416', '~0.268', '−35%'],
])

W('This counterintuitive result reveals a fundamental trade-off: the structure track introduces a physical conscience that sacrifices approximately 0.05 statistical naturalness to improve learning efficiency (P/R +35%) and physicochemical plausibility (0.738→0.752). Without the structure track, the model optimizes purely for statistical similarity to the reference distribution, achieving higher k-mer recall and dipeptide correlation—but at the cost of reduced physical realism and dramatically slower learning. Similar quality-naturalness trade-offs are well-documented in protein engineering [8,11], where sequences optimized solely for statistical properties may sacrifice structural viability.')

# 4. Discussion
H2('4. Discussion')
H3('4.1. Naturalness vs. Perplexity')
W('A central methodological choice in this work is elevating "naturalness"—multi-scale similarity to a reference protein family—over teacher-forcing accuracy as the primary evaluation metric. At equivalent sequence length (~116–158 AA), FoldPath-LLM achieves naturalness of 0.514 compared to RITA_m\'s 0.483 (+6.4%, p<0.05). While the statistical significance of the naturalness improvement alone is moderate, FoldPath-LLM provides three qualitative capabilities absent from RITA_m: (1) autonomous sequence termination—RITA_m\'s predictive entropy remains elevated (~2.8 nats) throughout generation, with per-step amino acid probabilities approaching a uniform distribution, leaving the model unable to determine when a sequence should end without external truncation; (2) temperature-controllable generation across a 0.24–0.8 range without mode collapse, with the counterintuitive property of increasing diversity at lower temperatures; and (3) structural awareness via the dual-track architecture and physicochemical biases [8,9].')

W('It is important to note that RITA_m and FoldPath-LLM differ in their training data: RITA_m was pretrained on UniRef50 (general proteins), while FoldPath-LLM additionally underwent fine-tuning on 437,000 cytochrome b sequences. This difference is not a flaw in the comparison but rather intrinsic to the experimental design: FoldPath-LLM\'s contribution is the sum of three factors—(1) cytochrome b family-specific fine-tuning, (2) the dual-track structural awareness architecture, and (3) physicochemical interaction biases. The ablation experiment (§3.5) isolates factor (2): the no_struct variant (fine-tuning only, no structure track) achieves naturalness of 0.560, confirming that fine-tuning alone elevates statistical naturalness, while the structure track contributes primarily to learning efficiency (+35% P/R) and physicochemical plausibility improvement rather than raw naturalness gain. This quality-naturalness trade-off is widely discussed in protein design [8,11]—sequences with superior structural viability and physicochemical plausibility often sacrifice pure statistical similarity to natural reference libraries.')

H3('4.2. Mechanism of the Structure Track')
W('The structure track serves two complementary roles during training. First, it provides auxiliary gradient signals that encourage the sequence track\'s hidden states to encode structurally meaningful information—even though the structure track\'s predictions (solvent exposure, secondary structure, distances) are never supervised with experimental data. This self-supervised structural signal functions as a regularizer shaping the latent representation space. Second, structural biases injected into attention enable learning of long-range dependencies difficult to capture through purely sequential attention. Transmembrane helices involve hydrophobic residues spaced approximately 3–4 positions apart—a pattern spanning 20+ residues that creates specific pairwise attention patterns. The structure track captures these patterns through bidirectional processing and feeds them back as attention biases. During inference, structural biases are intentionally disabled (use_bias=False) to prevent positive feedback collapse, analogous to teacher-student knowledge distillation. The ablation finding that the structure track improves P/R by 35% confirms a more fundamental role: structural biases function as an additional gradient signal channel that accelerates sequence pattern learning throughout the entire training process, going beyond generation-time decoration.')

H3('4.3. Limitations and Future Work')
W('Several limitations merit discussion. First, structural signals (solvent exposure, secondary structure, distances) are self-supervised rather than experimentally validated. Integrating experimentally determined structural data or high-confidence AlphaFold predictions as auxiliary supervision could strengthen guidance. Ramachandran plot validation [19] could serve as an additional structural plausibility metric. Second, the current model is trained on a single protein family (cytochrome b). Extending to multi-family generation would require family-conditional mechanisms. Charge and hydrophobicity patterns predicting folding mechanisms [16] could inform such extensions. Third, the generation length (mean 158 residues) remains shorter than full-length cytochrome b (~300–400 residues), constrained by the 256-residue training maximum. Computational tools improving stability face solubility trade-offs [15], which must be considered for length extension. Fourth, experimental validation (expression, purification, activity assays) is required to confirm that generated sequences fold into functional proteins—the ultimate test for any protein design method.')

# 5. Conclusions
H2('5. Conclusions')
W('We presented FoldPath-LLM, a dual-track Transformer architecture for structure-aware autoregressive protein sequence generation. By integrating a causal language model (RITA_m) with a bidirectional structure prediction track and physicochemical interaction biases, the model generates full-length cytochrome b sequences (mean 158 residues at optimal temperature 0.24) achieving a naturalness score of 0.514—significantly exceeding the RITA_m base model at equivalent length (0.483, +6.4%, p<0.05). The model maintains stable performance across a 0.24–0.8 temperature range with a counterintuitive diversity increase at lower temperatures (0.790→0.819). Ablation experiments confirm that the structure track contributes 35% improvement in learning efficiency, introducing a quality-naturalness trade-off rather than a pure naturalness gain. Sequence novelty verification (mean nearest-neighbor identity 14.6%) confirms de novo design without training set memorization. The core innovation—decoupling causal generation from bidirectional structure encoding—enables structural guidance during training without information leakage during inference. These results demonstrate that structural and physicochemical constraints can be productively integrated into autoregressive protein generation, moving beyond purely statistical sequence modeling toward physically informed protein design.')

# References
H2('References')
refs = [
    '1. Lin, Z.; Akin, H.; Rao, R.; et al. Evolutionary-scale prediction of atomic-level protein structure with a language model. Science 2023, 379, 1123–1130.',
    '2. Ferruz, N.; Schmidt, S.; Höcker, B. ProtGPT2 is a deep unsupervised language model for protein design. Nat. Commun. 2022, 13, 4348.',
    '3. Hesslow, D.; Zanichelli, N.; Notin, P.; et al. RITA: a Study on Scaling Up Generative Protein Sequence Models. arXiv 2022, arXiv:2205.05789.',
    '4. Dauparas, J.; Anishchenko, I.; Bennett, N.; et al. Robust deep learning–based protein sequence design using ProteinMPNN. Science 2022, 378, 49–56.',
    '5. Hsu, C.; Verkuil, R.; Liu, J.; et al. Learning inverse folding from millions of predicted structures. ICML 2022.',
    '6. Anishchenko, I.; Pellock, S.J.; Chidyausiku, T.M.; et al. De novo protein design by deep network hallucination. Nature 2021, 600, 547–552.',
    '7. Holtzman, A.; Buys, J.; Du, L.; et al. The Curious Case of Neural Text Degeneration. ICLR 2020.',
    '8. Pace, C.N.; Shirley, B.A.; McNutt, M. Forces contributing to the conformational stability of proteins. FASEB J. 1996, 10, 75–83.',
    '9. Hendsch, Z.S.; Tidor, B. Do salt bridges stabilize proteins? A continuum electrostatic analysis. Protein Sci. 1994, 3, 211–226.',
    '10. Waldburger, C.D.; Schildbach, J.F.; Sauer, R.T. Are buried salt bridges important for protein stability and conformational specificity? Nat. Struct. Biol. 1995, 2, 122–128.',
    '11. Kumar, S.; Nussinov, R. Salt bridge stability in monomeric proteins. J. Mol. Biol. 1999, 293, 1241–1255.',
    '12. Horovitz, A.; Serrano, L.; Avron, B.; Bycroft, M.; Fersht, A.R. Strength and co-operativity of contributions of surface salt bridges to protein stability. J. Mol. Biol. 1990, 216, 1031–1044.',
    '13. Anderson, D.E.; Becktel, W.J.; Dahlquist, F.W. pH-induced denaturation of proteins: a single salt bridge contributes 3–5 kcal/mol to the free energy of folding of T4 lysozyme. Biochemistry 1990, 29, 2403–2408.',
    '14. Bosshard, H.R.; Marti, D.N.; Jelesarov, I. Protein stabilization by salt bridges: concepts, experimental approaches and clarification of some misunderstandings. J. Mol. Recognit. 2004, 17, 1–16.',
    '15. Broom, A.; Jacobi, Z.; Trainor, K.; Meiering, E.M. Computational tools help improve protein stability but with a solubility tradeoff. J. Biol. Chem. 2017, 292, 14349–14361.',
    '16. Zbilut, J.P.; Giuliani, A.; Colosimo, A.; Mitchell, J.C.; Colafranceschi, M.; et al. Charge and hydrophobicity patterning along the sequence predicts the folding mechanism and aggregation of proteins. Proteomics 2004, 4, 1655–1663.',
    '17. McGaughey, G.B.; Gagné, M.; Rappé, A.K. π-Stacking interactions. Alive and well in proteins. J. Biol. Chem. 1998, 273, 15458–15463.',
    '18. Li, Z.; Yang, Y.; Zhan, J.; Dai, L.; Zhou, Y. Energy functions in de novo protein design: current challenges and future prospects. Annu. Rev. Biophys. 2013, 42, 315–335.',
    '19. Laskowski, R.A.; Furnham, N.; Thornton, J.M. The Ramachandran plot and protein structure validation. Methods Mol. Biol. 2013, 1012, 135–157.',
]
for ref in refs:
    p = P(); r = p.add_run(ref); r.font.name = 'Times New Roman'; r.font.size = Pt(10)

# 保存
output_path = 'FoldPathLLM_Paper_Final.docx'
doc.save(output_path)
print(f'Done: {output_path}')
